"""
通用文档读取模块。
支持格式：PDF, DOCX, DOC, XLSX, XLS, TXT, MD, HTML, RTF, CSV
"""
import os
import logging
from typing import Tuple, Optional

logger = logging.getLogger(__name__)

# 支持的文件扩展名
SUPPORTED_EXTENSIONS = {
    '.pdf', '.docx', '.doc', '.xlsx', '.xls',
    '.txt', '.md', '.markdown', '.html', '.htm',
    '.rtf', '.csv', '.json', '.yaml', '.yml', '.xml'
}


def is_supported(file_path: str) -> bool:
    """检查文件格式是否支持。"""
    ext = os.path.splitext(file_path)[1].lower()
    return ext in SUPPORTED_EXTENSIONS


def read_document(file_path: str) -> Tuple[str, bool]:
    """
    读取文档内容，返回 (文字内容, 是否为PDF)。
    
    PDF 需要特殊处理（混合识别），因此单独标记。
    其他格式直接提取纯文本返回。
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.pdf':
        return "", True  # PDF 走混合识别流程
    
    try:
        text = _extract_text(file_path, ext)
        if text:
            logger.info(f"📄 {ext} 文件读取成功，提取到 {len(text)} 字符")
        else:
            logger.warning(f"⚠️ {ext} 文件内容为空")
        return text or "", False
    except Exception as e:
        logger.error(f"❌ 文件读取失败 ({ext}): {e}")
        return "", False


def _extract_text(file_path: str, ext: str) -> Optional[str]:
    """根据文件类型提取文本。"""
    
    if ext in ('.txt', '.md', '.markdown'):
        return _read_text_file(file_path)
    
    elif ext == '.docx':
        return _read_docx(file_path)
    
    elif ext in ('.xlsx', '.xls'):
        return _read_excel(file_path)
    
    elif ext in ('.html', '.htm'):
        return _read_html(file_path)
    
    elif ext == '.csv':
        return _read_csv(file_path)
    
    elif ext == '.rtf':
        return _read_rtf(file_path)
    
    elif ext == '.doc':
        return _read_doc(file_path)
    
    elif ext in ('.json', '.yaml', '.yml', '.xml'):
        return _read_text_file(file_path)

    else:
        # 尝试当纯文本读取
        return _read_text_file(file_path)


def _read_text_file(file_path: str) -> str:
    """读取纯文本文件。"""
    encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
    for enc in encodings:
        try:
            with open(file_path, 'r', encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise ValueError(f"无法识别文件编码: {file_path}")


def _read_docx(file_path: str) -> str:
    """读取 Word (.docx) 文件，提取段落和表格内容。"""
    from docx import Document
    
    doc = Document(file_path)
    parts = []
    
    # 提取段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # 保留标题层级
            if para.style and para.style.name.startswith('Heading'):
                level = para.style.name.replace('Heading ', '').replace('Heading', '1')
                try:
                    level_num = int(level)
                except ValueError:
                    level_num = 1
                parts.append(f"{'#' * level_num} {text}")
            else:
                parts.append(text)
    
    # 提取表格
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(" | ".join(cells))
        if table_rows:
            # 生成 markdown 表格格式
            parts.append("")
            parts.append(table_rows[0])
            parts.append(" | ".join(["---"] * len(table.rows[0].cells)))
            parts.extend(table_rows[1:])
            parts.append("")
    
    return "\n".join(parts)


def _read_excel(file_path: str) -> str:
    """读取 Excel (.xlsx/.xls) 文件。"""
    from openpyxl import load_workbook
    
    wb = load_workbook(file_path, read_only=True, data_only=True)
    parts = []
    
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"## 工作表：{sheet_name}\n")
        
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(cell) if cell is not None else "" for cell in row]
            # 跳过全空行
            if any(c.strip() for c in cells):
                rows.append(" | ".join(cells))
        
        if rows:
            parts.append(rows[0])
            parts.append(" | ".join(["---"] * len(rows[0].split("|"))))
            parts.extend(rows[1:])
        parts.append("")
    
    wb.close()
    return "\n".join(parts)


def _read_html(file_path: str) -> str:
    """读取 HTML 文件，提取纯文本。"""
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        # 没有 bs4 时简单去标签
        content = _read_text_file(file_path)
        import re
        return re.sub(r'<[^>]+>', '', content)
    
    content = _read_text_file(file_path)
    soup = BeautifulSoup(content, 'html.parser')
    
    # 去掉 script 和 style
    for tag in soup(['script', 'style']):
        tag.decompose()
    
    return soup.get_text(separator='\n', strip=True)


def _read_csv(file_path: str) -> str:
    """读取 CSV 文件。"""
    import csv
    
    content = _read_text_file(file_path)
    lines = content.strip().split('\n')
    
    reader = csv.reader(lines)
    rows = [" | ".join(row) for row in reader]
    
    if len(rows) >= 2:
        # 插入 markdown 表头分隔符
        rows.insert(1, " | ".join(["---"] * len(rows[0].split("|"))))
    
    return "\n".join(rows)


def _read_rtf(file_path: str) -> str:
    """读取 RTF 文件（简单去标记提取文本）。"""
    import re
    content = _read_text_file(file_path)
    # 简易 RTF 文本提取
    text = re.sub(r'\\[a-z]+\d*\s?', '', content)
    text = re.sub(r'[{}]', '', text)
    return text.strip()


def _read_doc(file_path: str) -> str:
    """读取旧版 Word (.doc) 文件。"""
    # .doc 格式较复杂，尝试用 antiword 或提示用户转换
    import subprocess
    try:
        result = subprocess.run(
            ['textutil', '-convert', 'txt', '-stdout', file_path],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0:
            return result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    
    logger.warning("无法读取 .doc 文件，建议转换为 .docx 格式")
    return ""
